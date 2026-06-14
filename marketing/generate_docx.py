import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell padding/margins in twentieths of a point (dxa)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    """Set background color of a cell"""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def format_run_bold_italic(paragraph, text):
    """Processes bold/italic markdown notation inside a paragraph"""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            sub_text = part[2:-2]
            run = paragraph.add_run(sub_text)
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            sub_text = part[1:-1]
            run = paragraph.add_run(sub_text)
            run.italic = True
        else:
            paragraph.add_run(part)

def create_document():
    doc = Document()
    
    # Configure page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style Setup (Normal Style)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal

    md_path = 'project_proposal.md'
    if not os.path.exists(md_path):
        md_path = os.path.join('marketing', 'project_proposal.md')
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_headers = []
    table_rows = []

    for line in lines:
        stripped = line.strip()
        
        # Handle empty lines
        if not stripped:
            if in_table:
                # Process table
                build_table(doc, table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
            continue
            
        # Check if table row
        if stripped.startswith('|'):
            # Check if it's separator row
            if '---' in stripped and not any(c.isalnum() for c in stripped):
                continue
            cols = [col.strip() for col in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_headers = cols
            else:
                table_rows.append(cols)
            continue
        elif in_table:
            # Table ended
            build_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []

        # Handle headings
        if stripped.startswith('# '):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(12)
            run = h.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Navy Blue
            
        elif stripped.startswith('## '):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(8)
            run = h.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB) # Royal Blue
            
        elif stripped.startswith('### '):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6) # Light Blue

        elif stripped.startswith('#### '):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(stripped[5:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
            
        # Horizontal Rule
        elif stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p_border = OxmlElement('w:pBdr')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6')
            bottom_border.set(qn('w:space'), '1')
            bottom_border.set(qn('w:color'), 'CCCCCC')
            p_border.append(bottom_border)
            p._p.get_or_add_pPr().append(p_border)

        # Handle Bullet list items
        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            # Find indentation levels (represented by leading spaces)
            lead_spaces = len(line) - len(line.lstrip())
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (lead_spaces // 2))
            format_run_bold_italic(p, stripped[2:])
            
        elif stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            lead_spaces = len(line) - len(line.lstrip())
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (lead_spaces // 2))
            format_run_bold_italic(p, stripped[2:])

        # Numbered list items
        elif re.match(r'^\d+\.\s', stripped):
            match = re.match(r'^(\d+\.\s)(.*)', stripped)
            num_part = match.group(1)
            content_part = match.group(2)
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            num_run = p.add_run(num_part)
            num_run.bold = True
            
            format_run_bold_italic(p, content_part)

        # Standard Paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            format_run_bold_italic(p, stripped)

    if in_table:
        build_table(doc, table_headers, table_rows)

    out_path = 'project_proposal.docx'
    if 'marketing' in os.listdir('.'):
        out_path = os.path.join('marketing', 'project_proposal.docx')
    doc.save(out_path)
    print(f"Word document saved successfully to: {out_path}")


def build_table(doc, headers, rows):
    """Helper to construct a nicely formatted Word table"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.autofit = False
    
    # Custom widths
    widths = [Inches(1.8), Inches(3.2), Inches(1.5)] if len(headers) == 3 else [Inches(3.25), Inches(3.25)]
    
    # Format headers
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = widths[i] if i < len(widths) else Inches(2.0)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], '1E3A8A') # Navy
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=180, right=180)
        
    # Format rows
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = 'F3F4F6' if r_idx % 2 == 0 else 'FFFFFF' # Zebra striping
        
        for c_idx, val in enumerate(row):
            row_cells[c_idx].width = widths[c_idx] if c_idx < len(widths) else Inches(2.0)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            format_run_bold_italic(p, val)
            
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            
    # Add space after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)

if __name__ == '__main__':
    create_document()
